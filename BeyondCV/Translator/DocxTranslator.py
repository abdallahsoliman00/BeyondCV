from typing import Any, override

from docx import Document
from docx.document import Document as DocType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement  # pyright: ignore[reportUnknownVariableType]
from docx.oxml.ns import qn
from docx.oxml.table import CT_TcPr
from docx.shared import Cm, Pt
from docx.table import _Cell as DocxCell, Table as DocxTable  # pyright: ignore[reportPrivateUsage]
from docx.text.paragraph import Paragraph as DocxParagraph

from colour import Color

from BeyondCV.Translator.DocTranslator import DocTranslator
from BeyondCV.TableBuilder.Table import Table as CVTable, Row, Cell, Paragraph, Column
from BeyondCV.config import bcv_config as cfg


class DocxTranslator(DocTranslator):
    _HALIGN_MAP: dict[str, WD_ALIGN_PARAGRAPH] = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    @override
    def build(self, data: dict[str, Any]) -> str:
        tables = self._template.build(data)
        doc: DocType = Document()

        for section in doc.sections:
            section.top_margin = Cm(float(cfg.margin_top_cm))           # pyright: ignore[reportUnknownMemberType, reportUnknownArgumentType]
            section.bottom_margin = Cm(float(cfg.margin_bottom_cm))     # pyright: ignore[reportUnknownMemberType, reportUnknownArgumentType]
            section.left_margin = Cm(float(cfg.margin_left_cm))         # pyright: ignore[reportUnknownMemberType, reportUnknownArgumentType]
            section.right_margin = Cm(float(cfg.margin_right_cm))       # pyright: ignore[reportUnknownMemberType, reportUnknownArgumentType]

        for i, table_model in enumerate(tables):
            if i > 0:
                _ = doc.add_paragraph()
            self._add_table(doc, table_model)

        doc.save(str(self._doc_location))
        return str(self._doc_location)

    # ------------------------------------------------------------------ #
    #  Table rendering
    # ------------------------------------------------------------------ #

    def _add_table(self, doc: DocType, table_model: CVTable):
        if CVTable.are_columns(table_model.content):
            self._add_column_table(doc, table_model)
        else:
            self._add_row_table(doc, table_model)

    def _add_row_table(self, doc: DocType, table_model: CVTable):
        rows = [r for r in table_model.content if isinstance(r, Row)]
        if not rows:
            return

        max_cols = max(len(r.cells) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        self._remove_table_borders(table)

        for row_idx, row_model in enumerate(rows):
            table.rows[row_idx].height = Cm(row_model.min_height_cm)

            for col_idx, cell_model in enumerate(row_model.cells):
                cell = table.cell(row_idx, col_idx)
                self._fill_cell(cell, cell_model)

            if len(row_model.cells) < max_cols:
                start = table.cell(row_idx, len(row_model.cells) - 1)
                end = table.cell(row_idx, max_cols - 1)
                _ = start.merge(end)

    def _add_column_table(self, doc: DocType, table_model: CVTable):
        columns = [c for c in table_model.content if isinstance(c, Column)]
        if not columns:
            return

        col_count = len(columns)
        row_count = max(len(c.cells) for c in columns)
        if row_count == 0:
            return

        table = doc.add_table(rows=row_count, cols=col_count)
        self._remove_table_borders(table)

        for col_idx, col_model in enumerate(columns):
            for row_idx, cell_model in enumerate(col_model.cells):
                cell = table.cell(row_idx, col_idx)
                self._fill_cell(cell, cell_model)

    # ------------------------------------------------------------------ #
    #  Cell / Paragraph formatting
    # ------------------------------------------------------------------ #

    def _fill_cell(self, cell: DocxCell, cell_model: Cell):
        config = cell_model.config

        if config.width_cm > 0.0:
            cell.width = Cm(config.width_cm)

        self._set_cell_vertical_alignment(cell, config.content_alignment.get("vertical", "center"))

        self._remove_cell_shading(cell)
        if config.color is not None:
            self._set_cell_shading(cell, config.color)

        self._remove_cell_borders(cell)
        if config.show_borders:
            self._set_cell_borders(cell)

        for para_idx, para_model in enumerate(cell_model.paragraphs):
            p = cell.paragraphs[0] if para_idx == 0 else cell.add_paragraph()
            _ = p.clear()
            self._format_paragraph(p, para_model, config.content_alignment)

    def _format_paragraph(self, p: DocxParagraph, para_model: Paragraph, alignment: dict[str, str]):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        halign = alignment.get("horizontal", "left")
        p.alignment = self._HALIGN_MAP.get(halign, WD_ALIGN_PARAGRAPH.LEFT)

        if para_model.config.bullet:
            p.style = "List Bullet"

        run = p.add_run(para_model.text)
        run.font.name = para_model.config.font_name
        run.font.size = Pt(para_model.config.font_size_pt)
        run.bold = para_model.config.bold
        run.italic = para_model.config.italic
        run.underline = para_model.config.underline

    # ------------------------------------------------------------------ #
    #  XML helpers — vertical alignment
    # ------------------------------------------------------------------ #

    @staticmethod
    def _set_cell_vertical_alignment(cell: DocxCell, valign: str):
        tc = cell._tc  # pyright: ignore[reportPrivateUsage]
        tcPr: CT_TcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:vAlign"))
        if existing is not None:
            tcPr.remove(existing)
        elem = OxmlElement("w:vAlign")
        elem.set(qn("w:val"), valign)
        tcPr.append(elem)

    # ------------------------------------------------------------------ #
    #  XML helpers — shading / background colour
    # ------------------------------------------------------------------ #

    @staticmethod
    def _remove_cell_shading(cell: DocxCell):
        tc = cell._tc       # pyright: ignore[reportPrivateUsage]
        tcPr = tc.get_or_add_tcPr()
        for shd in tcPr.findall(qn("w:shd")):
            tcPr.remove(shd)

    @staticmethod
    def _set_cell_shading(cell: DocxCell, color: Color):
        hex_color = color.hex_l[1:]
        tc = cell._tc       # pyright: ignore[reportPrivateUsage]
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    # ------------------------------------------------------------------ #
    #  XML helpers — borders
    # ------------------------------------------------------------------ #

    @staticmethod
    def _remove_cell_borders(cell: DocxCell):
        tc = cell._tc       # pyright: ignore[reportPrivateUsage]
        tcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:tcBorders"))
        if existing is not None:
            tcPr.remove(existing)

    @staticmethod
    def _set_cell_borders(cell: DocxCell):
        tc = cell._tc       # pyright: ignore[reportPrivateUsage]
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            borders.append(border)
        tcPr.append(borders)

    @staticmethod
    def _remove_table_borders(table: DocxTable):
        tbl = table._tbl        # pyright: ignore[reportPrivateUsage]
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            borders.append(border)
        tblPr.append(borders)
